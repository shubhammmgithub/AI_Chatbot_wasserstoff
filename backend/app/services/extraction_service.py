import os
from typing import List, Dict, Any, Optional

import docx
import pytesseract
from PIL import Image, ImageOps
from pdf2image import convert_from_path
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.docstore.document import Document

from backend.app.core.logger import setup_logger
from backend.app.core.config import OCR_THRESHOLD, OCR_LANG, OCR_DPI, TESSERACT_CONFIG
from backend.app.utils.exceptions import FileExtractionError, ChunkingError

# Setup logger
logger = setup_logger("extraction_service")


class ExtractionService:
    """Service for extracting text from various file formats and chunking it."""

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean while preserving single newlines (better for chunking)."""
        if not text:
            return ""
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        lines = [" ".join(line.split()) for line in text.split("\n")]
        text = "\n".join(ln for ln in lines if ln.strip())
        return text.replace("\x00", "").replace("�", "").strip()

    @classmethod
    def handle_pdf(cls, file_path: str) -> List[Document]:
        """Native text via PyPDFLoader, OCR fallback if too little text."""
        try:
            logger.info(f"Processing PDF file: {file_path}")
            loader = PyPDFLoader(file_path)
            docs = loader.load()

            total = "".join(d.page_content or "" for d in docs).strip()
            if len(total) < OCR_THRESHOLD:
                logger.info(f"PDF has less than {OCR_THRESHOLD} characters, using OCR")
                # Scanned PDF → OCR every page image
                images = convert_from_path(file_path, dpi=OCR_DPI)
                out: List[Document] = []
                for i, img in enumerate(images):
                    logger.debug(f"OCR processing page {i+1}")
                    page_text = pytesseract.image_to_string(img, lang=OCR_LANG, config=TESSERACT_CONFIG)
                    cleaned = cls.clean_text(page_text)
                    if cleaned:
                        meta = {"source": os.path.basename(file_path), "page": i + 1}
                        out.append(Document(page_content=cleaned, metadata=meta))
                return out

            # Normalize metadata for native text: source and 1-based page
            fixed: List[Document] = []
            for d in docs:
                meta = dict(d.metadata or {})
                meta["source"] = os.path.basename(file_path)
                if meta.get("page") is not None:
                    try:
                        meta["page"] = int(meta["page"]) + 1
                    except Exception:
                        meta["page"] = 1
                else:
                    meta["page"] = 1
                fixed.append(Document(page_content=cls.clean_text(d.page_content), metadata=meta))
            return fixed
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {str(e)}")
            raise FileExtractionError(str(e), os.path.basename(file_path))

    @classmethod
    def handle_docx(cls, file_path: str) -> List[Document]:
        """Read paragraphs + tables, keep simple structure."""
        try:
            logger.info(f"Processing DOCX file: {file_path}")
            d = docx.Document(file_path)
            parts: List[str] = []
            parts.extend(p.text for p in d.paragraphs if p.text and p.text.strip())
            for table in d.tables:
                for row in table.rows:
                    parts.append("\t".join(cell.text.strip() for cell in row.cells))
            text = cls.clean_text("\n\n".join(parts))
            meta = {"source": os.path.basename(file_path)}
            return [Document(page_content=text, metadata=meta)] if text else []
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            raise FileExtractionError(str(e), os.path.basename(file_path))

    @classmethod
    def handle_txt(cls, file_path: str) -> List[Document]:
        """Process plain text files."""
        try:
            logger.info(f"Processing TXT file: {file_path}")
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = cls.clean_text(f.read())
            meta = {"source": os.path.basename(file_path)}
            return [Document(page_content=text, metadata=meta)] if text else []
        except Exception as e:
            logger.error(f"Error processing TXT file {file_path}: {str(e)}")
            raise FileExtractionError(str(e), os.path.basename(file_path))

    @classmethod
    def handle_image(cls, file_path: str) -> List[Document]:
        """OCR images; auto-fix orientation from EXIF."""
        try:
            logger.info(f"Processing image file: {file_path}")
            img = Image.open(file_path)
            img = ImageOps.exif_transpose(img)
            text = pytesseract.image_to_string(img, lang=OCR_LANG, config=TESSERACT_CONFIG)
            cleaned = cls.clean_text(text)
            if cleaned:
                meta = {"source": os.path.basename(file_path)}
                return [Document(page_content=cleaned, metadata=meta)]
            return []
        except Exception as e:
            logger.error(f"Error processing image file {file_path}: {str(e)}")
            raise FileExtractionError(str(e), os.path.basename(file_path))

    @classmethod
    def chunk_documents(cls, documents: List[Document], chunk_size: int = 1024, overlap: int = 200) -> List[Dict[str, Any]]:
        """Split into readable chunks with simple metadata (doc_id=filename)."""
        try:
            if not documents:
                logger.warning("No documents to chunk")
                return []

            # Basic guardrails so params are sane
            chunk_size = max(50, min(chunk_size, 2000))
            overlap = max(0, min(overlap, chunk_size // 2))
            
            logger.info(f"Chunking documents with size={chunk_size}, overlap={overlap}")

            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=overlap,
                length_function=len,
                separators=["\n\n", "\n", ". ", " ", ""],
            )

            chunks: List[Dict[str, Any]] = []
            for doc in documents:
                for i, chunk in enumerate(splitter.split_documents([doc])):
                    meta = {
                        "doc_id": doc.metadata.get("source", "unknown"),
                        "page": doc.metadata.get("page"),
                        "para": i + 1,
                        "text": chunk.page_content.strip(),
                    }
                    # Skip tiny fragments
                    if len(meta["text"]) < 50:
                        continue
                    chunks.append(meta)
            
            logger.info(f"Created {len(chunks)} chunks from {len(documents)} documents")
            return chunks
        except Exception as e:
            logger.error(f"Error chunking documents: {str(e)}")
            raise ChunkingError(str(e))

    @classmethod
    def extract_and_chunk(cls, file_path: str, chunk_size: int = 1024, overlap: int = 200) -> List[Dict[str, Any]]:
        """Extract text from a file and chunk it based on file type."""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            raise FileExtractionError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        logger.info(f"Processing file {file_path} with extension {ext}")

        if ext == ".pdf":
            docs = cls.handle_pdf(file_path)
        elif ext == ".docx":
            docs = cls.handle_docx(file_path)
        elif ext == ".txt":
            docs = cls.handle_txt(file_path)
        elif ext in [".png", ".jpg", ".jpeg"]:
            docs = cls.handle_image(file_path)
        else:
            logger.warning(f"Unsupported file type: {ext}")
            # Unsupported type → empty
            docs = []

        return cls.chunk_documents(docs, chunk_size=chunk_size, overlap=overlap)